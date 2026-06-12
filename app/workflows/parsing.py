from __future__ import annotations

import hashlib
import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING

from app.core.config import get_settings

if TYPE_CHECKING:
    from app.observability.tracing import Tracer
from app.core.errors import UnsupportedFileTypeError
from app.core.logging import get_logger
from app.locale import zh_CN as msg
from app.models.contracts import ParseStatus

log = get_logger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def normalize_text(text: str) -> str:
    return " ".join(text.split())


def text_hash(text: str) -> str:
    """Hash of normalized extracted text (not the original binary)."""
    return hashlib.sha256(normalize_text(text).encode("utf-8")).hexdigest()


def slugify(filename: str) -> str:
    stem = Path(filename).stem.lower()
    return re.sub(r"[^a-z0-9]+", "_", stem).strip("_") or "document"


@dataclass
class ParsedDocument:
    filename: str
    slug: str
    parse_status: ParseStatus
    text: str = ""
    page_texts: list[str] | None = None

    @property
    def document_hash(self) -> str | None:
        return text_hash(self.text) if self.text else None

    @property
    def char_count(self) -> int:
        return len(self.text)


def _parse_pdf_local(data: bytes) -> tuple[ParseStatus, str, list[str]]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        return "encrypted_pdf", "", []
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages).strip()
    if len(reader.pages) >= 1 and len(normalize_text(text)) < 30:
        return "scanned_pdf_requires_text_upload", "", []
    return "parsed", text, pages


def _parse_pdf_paddle(
    data: bytes,
    *,
    api_key: str,
    timeout: float,
    tracer: Tracer | None = None,
) -> tuple[ParseStatus, str, list[str]]:
    from app.workflows.paddle_ocr import call_paddleocr_vl, extract_text_from_response

    if tracer is not None:
        with tracer.span(
            "ocr_extract_pdf",
            input={"provider": "qianfan"},
        ) as ocr_span:
            payload = call_paddleocr_vl(data, api_key, timeout=timeout)
            text, page_texts = extract_text_from_response(payload)
            if len(normalize_text(text)) < 30:
                raise RuntimeError("paddle OCR 提取文本过短")
            ocr_span.update(output={"status": "parsed", "page_count": len(page_texts)})
            return "parsed", text, page_texts

    payload = call_paddleocr_vl(data, api_key, timeout=timeout)
    text, page_texts = extract_text_from_response(payload)
    if len(normalize_text(text)) < 30:
        raise RuntimeError("paddle OCR 提取文本过短")
    return "parsed", text, page_texts


def _parse_pdf(
    data: bytes,
    *,
    api_key: str | None = None,
    ocr_timeout: float = 120.0,
    tracer: Tracer | None = None,
) -> tuple[ParseStatus, str, list[str]]:
    if api_key:
        try:
            return _parse_pdf_paddle(
                data, api_key=api_key, timeout=ocr_timeout, tracer=tracer
            )
        except Exception as exc:
            log.warning("paddle OCR 失败，回退 pypdf 本地解析: %s", exc)
    return _parse_pdf_local(data)


def _parse_docx(data: bytes) -> tuple[ParseStatus, str]:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "parsed", "\n".join(parts).strip()


def parse_upload(
    filename: str,
    data: bytes,
    *,
    tracer: Tracer | None = None,
) -> ParsedDocument:
    slug = slugify(filename)
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise UnsupportedFileTypeError(msg.unsupported_file_type(filename))

    def _do_parse() -> ParsedDocument:
        try:
            ocr_used = False
            if suffix == ".pdf":
                settings = get_settings()
                ocr_used = bool(settings.qianfan_api_key)
                status, text, page_texts = _parse_pdf(
                    data,
                    api_key=settings.qianfan_api_key,
                    ocr_timeout=settings.paddle_ocr_timeout_seconds,
                    tracer=tracer,
                )
            elif suffix == ".docx":
                status, text = _parse_docx(data)
                page_texts = [text] if text else []
            else:
                status, text = "parsed", data.decode("utf-8", errors="replace").strip()
                page_texts = [text] if text else []
        except Exception:
            failed = ParsedDocument(filename=filename, slug=slug, parse_status="parse_failed")
            return failed, False

        if status == "parsed" and not text:
            status = "empty_text"
        return ParsedDocument(
            filename=filename,
            slug=slug,
            parse_status=status,
            text=text,
            page_texts=page_texts,
        ), ocr_used if suffix == ".pdf" else False

    if tracer is None:
        parsed, _ = _do_parse()
        return parsed

    with tracer.span(
        "parse_document",
        input={"filename": filename, "ext": suffix, "bytes_len": len(data)},
    ) as doc_span:
        parsed, ocr_used = _do_parse()
        doc_span.update(
            output={
                "parse_status": parsed.parse_status,
                "char_count": parsed.char_count,
                "document_hash": parsed.document_hash,
                "ocr_used": ocr_used,
            }
        )
        return parsed
